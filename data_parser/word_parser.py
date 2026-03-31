import os
import hashlib
import subprocess
import tempfile
from typing import List
from .base_parser import BaseParser, DataBlock, ModalityType


class WordParser(BaseParser):
    """Word文件解析器

    将Word文件解析为数据块，整个文档合并为一个数据块。
    """

    def __init__(self, cache_dir: str = "cache/data_blocks"):
        super().__init__()
        self.supported_extensions = ['docx', 'doc']
        self.cache_dir = cache_dir

    def _get_cache_path(self, file_path: str) -> str:
        """生成缓存目录路径"""
        file_hash = hashlib.md5(os.path.abspath(file_path).encode()).hexdigest()[:12]
        file_name = os.path.splitext(os.path.basename(file_path))[0]

        cache_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            self.cache_dir,
            f"{file_name}_{file_hash}"
        )

        os.makedirs(cache_path, exist_ok=True)
        return cache_path

    def parse(self, file_path: str) -> List[DataBlock]:
        """解析Word文件"""
        if not self.can_parse(file_path):
            raise ValueError(f"Unsupported file format: {file_path}")

        ext = file_path.lower().split('.')[-1]

        if ext == 'doc':
            return self._parse_doc(file_path)
        else:
            return self._parse_docx(file_path)

    def _parse_docx(self, file_path: str) -> List[DataBlock]:
        """解析docx文件 - 整个文档合并为一个数据块"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for Word file parsing. "
                "Install it with: pip install python-docx"
            )

        blocks = []
        try:
            doc = Document(file_path)
        except Exception as e:
            # 文件打开失败，返回空列表，让上层处理
            print(f"[WordParser] 无法打开Word文件: {e}")
            return blocks

        # 获取缓存目录
        cache_path = self._get_cache_path(file_path)

        # 收集所有内容
        all_texts = []

        # 解析段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_texts.append(paragraph.text.strip())

        # 解析表格
        for table in doc.tables:
            try:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                table_text = self._table_to_text(table_data)
                all_texts.append(table_text)
            except Exception as e:
                print(f"[WordParser] 解析表格失败: {e}")

        # 合并所有内容为一个数据块
        if all_texts:
            combined_text = "\n\n".join(all_texts)
            text_cache_path = os.path.join(cache_path, "document.txt")
            try:
                with open(text_cache_path, 'w', encoding='utf-8') as f:
                    f.write(combined_text)
            except Exception as e:
                print(f"[WordParser] 保存文本缓存失败: {e}")
                text_cache_path = None

            block = DataBlock(
                block_id=self._generate_block_id(file_path, 0),
                modality=ModalityType.TEXT,
                addr=text_cache_path,
                file_path=file_path,
                metadata={
                    'source': 'python-docx',
                    'content_count': len(all_texts),
                    'text_length': len(combined_text)
                }
            )
            blocks.append(block)

        return blocks

    def _parse_doc(self, file_path: str) -> List[DataBlock]:
        """解析doc文件（需要LibreOffice转换）"""
        temp_dir = tempfile.gettempdir()
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        docx_path = os.path.join(temp_dir, f"{base_name}.docx")

        try:
            # 尝试使用LibreOffice转换
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'docx',
                 '--outdir', temp_dir, file_path],
                capture_output=True,
                text=True,
                timeout=60  # 添加超时限制
            )

            if result.returncode == 0 and os.path.exists(docx_path):
                blocks = self._parse_docx(docx_path)
                try:
                    os.remove(docx_path)
                except Exception:
                    pass
                return blocks
            else:
                error_msg = result.stderr if result.stderr else "未知错误"
                print(f"[WordParser] LibreOffice转换失败: {error_msg}")
                # 返回空列表而不是抛出异常，让文件仍能被处理
                return []

        except subprocess.TimeoutExpired:
            print("[WordParser] LibreOffice转换超时")
            return []
        except FileNotFoundError:
            print("[WordParser] 未找到LibreOffice (soffice命令)，请安装LibreOffice或将.doc转换为.docx格式")
            return []
        except Exception as e:
            print(f"[WordParser] 解析.doc文件失败: {e}")
            return []

    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """将表格转换为文本"""
        rows = []
        for row in table_data:
            row_text = ' | '.join(str(cell) if cell else '' for cell in row)
            rows.append(row_text)
        return '\n'.join(rows)